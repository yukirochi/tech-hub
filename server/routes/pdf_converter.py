from fastapi import APIRouter, UploadFile, File, HTTPException, Response
from pydantic import BaseModel
import io
import hashlib
from cache.cached_image import get_cached_image, set_cached_image

router = APIRouter(
    tags=["general"],
)

# Try to import PDF libraries
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

@router.post("/pdf_to_word")
async def pdf_to_word(file: UploadFile = File(...)):
    """Convert PDF to Word document"""
    try:
        if not PYPDF2_AVAILABLE:
            return {
                "error": "PDF conversion not available",
                "message": "PyPDF2 library is not installed. Please install it with: pip install PyPDF2",
                "status": "unavailable"
            }
        
        if not DOCX_AVAILABLE:
            return {
                "error": "Word document creation not available",
                "message": "python-docx library is not installed. Please install it with: pip install python-docx",
                "status": "unavailable"
            }
        
        file_content = await file.read()
        
        if not file_content:
            raise HTTPException(status_code=400, detail="File is empty")
        
        # Check cache
        hash_key = hashlib.sha256(file_content).hexdigest() + "_pdf_to_word"
        cached = get_cached_image(hash_key)
        if cached:
            return Response(content=cached, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        # Read PDF
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Create Word document
        doc = Document()
        doc.add_heading('Converted from PDF', 0)
        
        # Extract text from each page
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text.strip():
                doc.add_heading(f'Page {page_num + 1}', level=1)
                doc.add_paragraph(text)
                doc.add_paragraph('')  # Add spacing
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        result_bytes = output.getvalue()
        
        # Cache result
        set_cached_image(hash_key, result_bytes)
        
        return Response(
            content=result_bytes,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="converted_{file.filename}.docx"'}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")

@router.post("/word_to_pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    """Convert Word document to PDF"""
    try:
        if not DOCX_AVAILABLE:
            return {
                "error": "Word document reading not available",
                "message": "python-docx library is not installed. Please install it with: pip install python-docx",
                "status": "unavailable"
            }
        
        if not REPORTLAB_AVAILABLE:
            return {
                "error": "PDF creation not available",
                "message": "reportlab library is not installed. Please install it with: pip install reportlab",
                "status": "unavailable"
            }
        
        file_content = await file.read()
        
        if not file_content:
            raise HTTPException(status_code=400, detail="File is empty")
        
        # Check cache
        hash_key = hashlib.sha256(file_content).hexdigest() + "_word_to_pdf"
        cached = get_cached_image(hash_key)
        if cached:
            return Response(content=cached, media_type='application/pdf')
        
        # Read Word document
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        # Create PDF
        output = io.BytesIO()
        pdf = SimpleDocTemplate(output, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Extract text from Word document
        for para in doc.paragraphs:
            if para.text.strip():
                p = Paragraph(para.text, styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 0.2 * inch))
        
        # Build PDF
        pdf.build(story)
        output.seek(0)
        result_bytes = output.getvalue()
        
        # Cache result
        set_cached_image(hash_key, result_bytes)
        
        return Response(
            content=result_bytes,
            media_type='application/pdf',
            headers={'Content-Disposition': f'attachment; filename="converted_{file.filename}.pdf"'}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")
